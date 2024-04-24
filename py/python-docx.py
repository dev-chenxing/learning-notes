from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

file_name = "Lorem Ipsum"
directory = ""

# create the document
doc = Document()                                            

# set margin to 2 cm
section = doc.sections[0]
section.left_margin, section.right_margin = Cm(2), Cm(2)    

normal_style = doc.styles["Normal"]
normal_font = normal_style.font
normal_font.name = "Helvetica"                  # set font family
normal_font.size = Pt(12)                       # set font size to 12pt
normal_paragraph_format = normal_style.paragraph_format   
normal_paragraph_format.line_spacing = Pt(14)   # set line spacing to 14pt

# add a centered paragraph
center_paragraph = doc.add_paragraph("Lorem ipsum")
center_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER      

# add a plain paragraph
doc.add_paragraph("Lorem ipsum dolor sit amet, ")           

# add a paragraph with color red
red_paragraph = doc.add_paragraph("consectetur adipiscing elit, ")
red_paragraph.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  

# add a paragraph with color blue
blue_paragraph = doc.add_paragraph(
    "sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. ")
blue_paragraph.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0xFF) 

# save the document to the output path
output_path = f"{directory}{file_name}.docx"
doc.save(output_path)